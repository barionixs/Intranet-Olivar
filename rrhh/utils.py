"""Cálculos que se derivan de otros datos en vez de guardarse aparte
(así nadie tiene que actualizarlos a mano): antigüedad y saldo de
días administrativos/vacaciones (cupo anual fijo, ver CUPO_DIAS_ANUAL)."""

import hashlib
import io
import logging
import subprocess
import tempfile
from datetime import date
from email.mime.image import MIMEImage
from pathlib import Path

import qrcode
from django.conf import settings
from django.core.files.base import ContentFile
from django.core.mail import EmailMultiAlternatives
from django.db.models import Sum
from django.template.defaultfilters import date as date_filter
from django.template.loader import render_to_string
from django.urls import reverse
from django.utils import timezone
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from docx.text.paragraph import Paragraph
from xhtml2pdf import pisa

from .models import CargoUnico, FichaLaboral, FirmaSolicitud, SolicitudPermiso

logger = logging.getLogger(__name__)


def calcular_antiguedad(fecha_ingreso, hoy=None):
    """Retorna (años, meses) completos entre fecha_ingreso y hoy."""
    hoy = hoy or date.today()
    años = hoy.year - fecha_ingreso.year
    meses = hoy.month - fecha_ingreso.month
    if hoy.day < fecha_ingreso.day:
        meses -= 1
    if meses < 0:
        años -= 1
        meses += 12
    return max(años, 0), max(meses, 0)


# Cupo anual fijo por funcionario, igual para todos independiente de
# antigüedad o tipo de contrato (reemplaza el cálculo por antigüedad que
# existía antes solo como estimación). "Sin goce" y "Otro" no tienen
# cupo — no son parte de la asignación fija, así que no se limitan acá.
CUPO_DIAS_ANUAL = {
    SolicitudPermiso.Tipo.ADMINISTRATIVO: 6,
    SolicitudPermiso.Tipo.VACACIONES: 15,
}


def dias_usados_en_año(funcionario, tipo, año):
    """Suma de días de solicitudes APROBADAS de ese tipo cuyo inicio
    cae en `año` — solo lo aprobado descuenta cupo; lo pendiente o
    rechazado no, así no hace falta "devolver" días si se rechaza."""
    total = SolicitudPermiso.objects.filter(
        funcionario=funcionario,
        tipo=tipo,
        estado=SolicitudPermiso.Estado.APROBADO,
        fecha_inicio__year=año,
    ).aggregate(total=Sum("dias_solicitados"))["total"]
    return total or 0


def dias_disponibles(funcionario, tipo, año=None):
    """Cupo restante de ese tipo para `año` (el actual si no se indica),
    o None si ese tipo no tiene cupo fijo (sin goce, otro) — sin límite
    que hacer cumplir en ese caso."""
    cupo = CUPO_DIAS_ANUAL.get(tipo)
    if cupo is None:
        return None
    año = año or date.today().year
    return cupo - dias_usados_en_año(funcionario, tipo, año)


def resumen_saldos(funcionario, año=None):
    """Saldo de cada tipo con cupo fijo, para mostrar en Mi Ficha. Se
    reinicia solo cada 1 de enero: no hay un contador que resetear a
    mano, el cupo restante siempre se calcula sobre las solicitudes
    aprobadas de `año`, así que un año nuevo parte automáticamente
    en cupo completo."""
    año = año or date.today().year
    etiquetas = dict(SolicitudPermiso.Tipo.choices)
    resumen = []
    for tipo, cupo in CUPO_DIAS_ANUAL.items():
        usados = dias_usados_en_año(funcionario, tipo, año)
        resumen.append({
            "tipo": tipo,
            "etiqueta": etiquetas[tipo],
            "cupo": cupo,
            "usados": usados,
            "disponibles": cupo - usados,
        })
    return resumen


# --- Firma electrónica simple (Ley 19.799) ----------------------------------
# Cada SolicitudPermiso lleva 3 firmas en orden: interesado, jefatura
# directa, Alcaldesa. No es firma avanzada ni usa certificador
# acreditado — la atribución razonable a la persona se logra pidiendo
# la contraseña de nuevo al firmar (no basta la sesión activa) + hash
# del contenido + IP + timestamp, registrado en FirmaSolicitud.

def resolver_jefatura_directa(funcionario):
    """Asignación directa, sin cadenas ni niveles de cargo: si esta
    persona tiene un `jefe_directo` asignado a mano (excepción, ej.
    Norma Vidal), ese es el firmante — sin seguir subiendo más allá.
    Si no, se usa el Jefe de su departamento (o el del departamento
    padre si el suyo no tiene uno propio, un solo salto por nivel,
    tal como está armado el organigrama). Devuelve None si nada de
    esto está configurado todavía — hay que bloquear el envío de la
    solicitud con un mensaje claro, no es un error real."""
    if funcionario.jefe_directo_id:
        return funcionario.jefe_directo
    depto = funcionario.departamento
    visitados = set()
    while depto and depto.id not in visitados:
        if depto.jefe_id and depto.jefe_id != funcionario.id:
            return depto.jefe
        visitados.add(depto.id)
        depto = depto.departamento_padre
    return None


def hash_solicitud(solicitud):
    """sha256 de los campos que definen QUÉ se está firmando. No incluye
    `estado` (cambia después de firmar) ni fecha_solicitud/fecha_resolucion
    (metadata, no contenido) — así la firma queda ligada a la versión
    exacta de la solicitud en ese momento."""
    campos = "|".join(str(v) for v in [
        solicitud.pk,
        solicitud.funcionario_id,
        solicitud.tipo,
        solicitud.fecha_inicio,
        solicitud.fecha_termino,
        solicitud.dias_solicitados,
        solicitud.motivo,
    ])
    return hashlib.sha256(campos.encode("utf-8")).hexdigest()


def obtener_ip_cliente(request):
    x_forwarded_for = request.META.get("HTTP_X_FORWARDED_FOR")
    if x_forwarded_for:
        return x_forwarded_for.split(",")[0].strip()
    return request.META.get("REMOTE_ADDR")


def hash_firma(firma):
    """sha256 de ESTA firma en particular, no solo del documento: liga
    el hash del contenido a quién firmó, cuándo y desde qué IP. Sin
    esto, las 3 firmas de una misma solicitud (interesado, jefatura,
    Alcaldesa) tendrían el mismo hash por firmar el mismo documento —
    correcto para probar que todos vieron la misma versión, pero no
    sirve para distinguir una firma de otra. Requiere que `fecha_firma`
    e `ip_firma` ya estén asignados en la instancia antes de llamarla."""
    campos = "|".join(str(v) for v in [
        hash_solicitud(firma.solicitud),
        firma.funcionario_firmante_id,
        firma.fecha_firma,
        firma.ip_firma,
    ])
    return hashlib.sha256(campos.encode("utf-8")).hexdigest()


def crear_firmas_para_solicitud(solicitud):
    """Genera las 3 FirmaSolicitud pendientes al crear una solicitud.
    El cargo único (Alcaldesa) puede no tener funcionario asignado
    todavía — esa firma queda pendiente sin firmante hasta que se
    configure, sin bloquear el resto."""
    alcaldesa = CargoUnico.objects.filter(nombre=CargoUnico.Nombre.ALCALDESA).first()
    jefatura_directa = resolver_jefatura_directa(solicitud.funcionario)

    pasos = [
        (1, FirmaSolicitud.RolFirmante.INTERESADO, solicitud.funcionario),
        (2, FirmaSolicitud.RolFirmante.JEFATURA_DIRECTA, jefatura_directa),
        (3, FirmaSolicitud.RolFirmante.ALCALDESA, alcaldesa.funcionario if alcaldesa else None),
    ]
    FirmaSolicitud.objects.bulk_create([
        FirmaSolicitud(solicitud=solicitud, orden=orden, rol_firmante=rol, funcionario_firmante=firmante)
        for orden, rol, firmante in pasos
    ])


def generar_comprobante_firma(solicitud):
    """PDF de respaldo/auditoría con el detalle de las 4 firmas (no es
    la firma en sí, esa ya quedó registrada en cada FirmaSolicitud).
    Se genera una sola vez, al momento exacto en que la solicitud
    queda aprobada."""
    firmas = list(
        solicitud.firmas.select_related("funcionario_firmante").order_by("orden")
    )
    html = render_to_string(
        "rrhh/comprobante_firma.html", {"solicitud": solicitud, "firmas": firmas}
    )
    buffer = io.BytesIO()
    resultado = pisa.CreatePDF(src=html, dest=buffer)
    if resultado.err:
        return
    solicitud.comprobante_firma.save(
        f"solicitud_{solicitud.pk}.pdf", ContentFile(buffer.getvalue()), save=True
    )


# --- Formato institucional (Solicitud de Permiso Administrativo) ------------
# Se rellena EL MISMO archivo Word institucional (plantillas_word/PERMISOS
# ADMINISTRATIVOS.docx: logos, escudo y pie de página incluidos) en vez de
# recrear su diseño en HTML — así el resultado es idéntico al formato
# oficial, solo con los campos completados. En cada campo de firma del
# documento se agrega el detalle de la firma electrónica simple de esa
# persona (nombre, RUT, fecha/hora y hash) en vez de una firma manuscrita.
# El .docx rellenado se convierte a PDF con LibreOffice en modo headless
# para poder verlo/descargarlo como los demás documentos del sistema — en
# local, con el LibreOffice instalado en la máquina; en Vercel, mandando
# la conversión a Gotenberg (ver convertir_docx_a_pdf / GOTENBERG_URL).

PLANTILLA_PERMISO_ADMINISTRATIVO = settings.BASE_DIR / "plantillas_word" / "PERMISOS ADMINISTRATIVOS.docx"

TIPOS_CON_FORMATO_INSTITUCIONAL = {
    SolicitudPermiso.Tipo.ADMINISTRATIVO,
    SolicitudPermiso.Tipo.SIN_GOCE,
    SolicitudPermiso.Tipo.VACACIONES,
}

# Fila (0-indexada) de la tabla de tipos de permiso en la plantilla que
# corresponde a cada `Tipo` de SolicitudPermiso.
_FILA_TIPO_EN_PLANTILLA = {
    SolicitudPermiso.Tipo.ADMINISTRATIVO: 0,  # Permiso con sueldo (Art. 108)
    SolicitudPermiso.Tipo.SIN_GOCE: 1,        # Permiso sin sueldo (Art. 109)
    SolicitudPermiso.Tipo.VACACIONES: 2,      # Feriado legal (Arts. 101-103)
}

# Texto exacto (una vez sin espacios de relleno) de cada etiqueta de firma
# en la plantilla, con el rol de FirmaSolicitud que le corresponde y el
# margen izquierdo/derecho del bloque que se agrega debajo — medido
# directamente sobre la línea "____" de esa firma en el documento
# original (plantillas_word/PERMISOS ADMINISTRATIVOS.docx renderizado),
# no estimado, para que el bloque quede contenido en el mismo ancho que
# ocupa esa línea en vez de extenderse por todo el ancho de la página.
_ETIQUETAS_FIRMA_EN_PLANTILLA = {
    "FIRMA INTERESADO (A)": (FirmaSolicitud.RolFirmante.INTERESADO, Pt(288), Pt(16)),
    "V° B° JEFE DIRECTO": (FirmaSolicitud.RolFirmante.JEFATURA_DIRECTA, Pt(0), Pt(323)),
    "A L C A L D E S A": (FirmaSolicitud.RolFirmante.ALCALDESA, Pt(288), Pt(16)),
}


def tiene_formato_institucional(solicitud):
    """El formulario en papel solo cubre permiso con goce (Art. 108),
    sin goce (Art. 109) y feriado legal — "Otro" no tiene un campo al
    que mapear, así que ese tipo no ofrece este formato."""
    return solicitud.tipo in TIPOS_CON_FORMATO_INSTITUCIONAL


def _fecha_en_letras(fecha):
    return date_filter(fecha, "d \\d\\e F \\d\\e Y")


def _texto_firma_electronica(firma):
    """Líneas que se imprimen bajo cada campo de firma del documento,
    o el estado pendiente/rechazado si aún no se ha resuelto.

    Deliberadamente NO incluye el RUT ni el hash de la firma — solo
    nombre, cargo, fecha/hora e institución — para no imprimir datos
    sensibles en un documento que puede circular en papel; el RUT y el
    hash siguen disponibles para quien corresponda en el comprobante de
    firma y en la verificación pública por código."""
    if firma is None or firma.estado == FirmaSolicitud.Estado.PENDIENTE:
        if firma is not None and firma.funcionario_firmante:
            return [f"Pendiente de firma de {firma.funcionario_firmante.nombre_completo}"]
        return ["Pendiente — sin firmante asignado"]
    if firma.estado == FirmaSolicitud.Estado.RECHAZADO:
        texto = (
            f"Rechazado por {firma.funcionario_firmante.nombre_completo} "
            f"el {date_filter(firma.fecha_firma, 'd/m/Y H:i')}"
        )
        if firma.comentario:
            texto += f' — "{firma.comentario}"'
        return [texto]
    return [
        f"Firmado por: {firma.funcionario_firmante.nombre_completo}",
        firma.funcionario_firmante.get_cargo_display(),
        date_filter(firma.fecha_firma, "d/m/Y H:i"),
        "Ilustre Municipalidad de Olivar",
    ]


def _completar_campo(cell, valor):
    """Escribe `valor` al final del último run de la celda, para
    conservar el formato (negrita, fuente) que ya trae la plantilla en
    ese campo en vez de reemplazarlo por texto con formato por defecto."""
    if valor is None or valor == "":
        return
    parrafo = cell.paragraphs[0]
    if parrafo.runs:
        parrafo.runs[-1].text += str(valor)
    else:
        parrafo.add_run(str(valor))


def _insertar_parrafo_antes(parrafo, lineas, indent_izquierdo, indent_derecho):
    """python-docx no trae `insert_paragraph_before`: se arma el <w:p> a
    mano y se inserta antes de la línea "____" de esa firma — como una
    firma manuscrita real, que va ARRIBA de la línea, no debajo de la
    etiqueta. Los márgenes izquierdo/derecho confinan el bloque al mismo
    ancho que ocupa esa línea, en vez de todo el ancho de la página.
    `lineas` es una lista: cada elemento va en su propia línea dentro
    del mismo párrafo (nombre, cargo, fecha, institución por separado)."""
    nuevo_p = OxmlElement("w:p")
    parrafo._p.addprevious(nuevo_p)
    nuevo_parrafo = Paragraph(nuevo_p, parrafo._parent)
    nuevo_parrafo.paragraph_format.left_indent = indent_izquierdo
    nuevo_parrafo.paragraph_format.right_indent = indent_derecho
    for i, linea in enumerate(lineas):
        run = nuevo_parrafo.add_run(linea)
        run.italic = True
        run.font.size = Pt(8)
        if i < len(lineas) - 1:
            run.add_break()
    return nuevo_parrafo


def _generar_qr(datos):
    qr = qrcode.QRCode(border=1, box_size=6)
    qr.add_data(datos)
    qr.make(fit=True)
    imagen = qr.make_image(fill_color="black", back_color="white")
    buffer = io.BytesIO()
    imagen.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _quitar_bordes_tabla(tabla):
    tbl_pr = tabla._tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "nil")
        bordes.append(el)
    tbl_pr.append(bordes)


def _agregar_pie_verificacion(documento, solicitud):
    """Agrega al pie de página (sin tocar el pie institucional que ya
    trae la plantilla) un QR que enlaza a la verificación pública de
    esta solicitud — mismo código/vista que ya usa el sistema para
    "Verificar firma" — junto a una glosa breve de la Ley 19.799, igual
    que en el comprobante de firma."""
    url_verificacion = settings.SITE_URL.rstrip("/") + reverse(
        "verificar_firma", args=[solicitud.codigo_verificacion]
    )
    qr = _generar_qr(url_verificacion)

    footer = documento.sections[0].footer
    tabla = footer.add_table(rows=1, cols=2, width=Inches(6.0))
    _quitar_bordes_tabla(tabla)
    tabla.columns[0].width = Inches(0.9)
    tabla.columns[1].width = Inches(5.1)

    celda_qr = tabla.rows[0].cells[0]
    celda_qr.paragraphs[0].add_run().add_picture(qr, width=Inches(0.8))

    celda_texto = tabla.rows[0].cells[1]
    run = celda_texto.paragraphs[0].add_run(
        "Firma electrónica simple conforme a la Ley N.º 19.799 sobre documentos electrónicos, firma "
        "electrónica y servicios de certificación de dicha firma. Escanee el código para verificar la "
        f"autenticidad de este documento y el estado de sus firmas, o ingrese el código "
        f"{solicitud.codigo_verificacion} en la Intranet, sección \"Verificar firma\"."
    )
    run.italic = True
    run.font.size = Pt(6.5)


def rellenar_documento_permiso(solicitud):
    """Devuelve los bytes del .docx institucional con los datos de la
    solicitud y el estado ACTUAL de sus firmas ya completados — útil
    tanto para la vista previa en cualquier momento como para la
    versión final guardada al aprobarse."""
    firmas = {
        firma.rol_firmante: firma
        for firma in solicitud.firmas.select_related("funcionario_firmante").all()
    }
    documento = Document(str(PLANTILLA_PERMISO_ADMINISTRATIVO))
    for seccion in documento.sections:
        seccion.page_width = Mm(216)
        seccion.page_height = Mm(330)
    fecha_doc = solicitud.fecha_resolucion or solicitud.fecha_solicitud

    for parrafo in documento.paragraphs:
        if parrafo.text.startswith("FECHA ") and parrafo.runs:
            parrafo.runs[0].text = f"FECHA {date_filter(fecha_doc, 'd/m/Y')}"
            for extra in parrafo.runs[1:]:
                extra.text = ""
        elif parrafo.text.startswith("OLIVAR,") and len(parrafo.runs) > 1:
            parrafo.runs[-1].text = f" {_fecha_en_letras(fecha_doc)}"

    tabla_datos = documento.tables[0]
    try:
        grado = solicitud.funcionario.ficha_laboral.grado
    except FichaLaboral.DoesNotExist:
        grado = ""
    _completar_campo(tabla_datos.rows[0].cells[1], solicitud.funcionario.nombre_completo)
    _completar_campo(tabla_datos.rows[1].cells[1], solicitud.funcionario.rut)
    _completar_campo(tabla_datos.rows[2].cells[1], grado)
    _completar_campo(
        tabla_datos.rows[3].cells[1],
        f"{solicitud.funcionario.get_cargo_display()} · {solicitud.funcionario.departamento.nombre}",
    )

    fila_tipo = _FILA_TIPO_EN_PLANTILLA.get(solicitud.tipo)
    if fila_tipo is not None:
        _completar_campo(documento.tables[1].rows[fila_tipo].cells[1], solicitud.dias_solicitados)

    tabla_fechas = documento.tables[2]
    _completar_campo(tabla_fechas.rows[0].cells[1], date_filter(solicitud.fecha_inicio, "d/m/Y"))
    _completar_campo(tabla_fechas.rows[1].cells[1], date_filter(solicitud.fecha_termino, "d/m/Y"))

    _completar_campo(documento.tables[3].rows[0].cells[1], solicitud.motivo)

    seccion = documento.sections[0]
    ancho_texto = seccion.page_width - seccion.left_margin - seccion.right_margin

    anterior = None
    id_marca_agua = 9000
    for parrafo in list(documento.paragraphs):
        info = _ETIQUETAS_FIRMA_EN_PLANTILLA.get(parrafo.text.strip())
        if info and anterior is not None:
            rol, indent_izq, indent_der = info
            firma = firmas.get(rol)
            nuevo_parrafo = _insertar_parrafo_antes(
                anterior, _texto_firma_electronica(firma), indent_izq, indent_der
            )
            if firma is not None and firma.estado == FirmaSolicitud.Estado.FIRMADO:
                id_marca_agua += 1
                ancho_columna = ancho_texto - indent_izq - indent_der
                _agregar_marca_agua(nuevo_parrafo, id_marca_agua, indent_izq, ancho_columna)
        anterior = parrafo

    _agregar_pie_verificacion(documento, solicitud)

    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


def _ruta_libreoffice():
    for candidato in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
    ):
        if candidato == "soffice" or Path(candidato).exists():
            return candidato
    return None


def _convertir_via_gotenberg(contenido_docx):
    """Convierte llamando a un Gotenberg (LibreOffice en Docker) desplegado
    aparte — es la vía que se usa en Vercel, donde no hay LibreOffice
    instalado. Se activa solo si GOTENBERG_URL está configurado."""
    import requests

    url = settings.GOTENBERG_URL.rstrip("/") + "/forms/libreoffice/convert/document"
    auth = None
    if settings.GOTENBERG_USER:
        auth = (settings.GOTENBERG_USER, settings.GOTENBERG_PASSWORD)
    try:
        respuesta = requests.post(
            url,
            files={"files": ("documento.docx", contenido_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            auth=auth,
            # El servicio gratuito puede estar "dormido" y tardar en despertar.
            timeout=90,
        )
    except requests.RequestException:
        return None
    if respuesta.status_code != 200:
        return None
    return respuesta.content


def convertir_docx_a_pdf(contenido_docx):
    """Convierte a PDF: usa Gotenberg (remoto) si GOTENBERG_URL está
    configurado, o LibreOffice local en modo headless si no — cada
    llamada local usa un perfil de usuario propio, así conversiones
    simultáneas (dos personas viendo la vista previa a la vez) no
    compiten por el mismo perfil de LibreOffice y se bloqueen entre sí."""
    if settings.GOTENBERG_URL:
        return _convertir_via_gotenberg(contenido_docx)

    soffice = _ruta_libreoffice()
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / "documento.docx"
        docx_path.write_bytes(contenido_docx)
        perfil = (tmp_path / "perfil_lo").as_posix()
        try:
            resultado = subprocess.run(
                [
                    soffice, "--headless", "--norestore",
                    f"-env:UserInstallation=file:///{perfil}",
                    "--convert-to", "pdf", "--outdir", str(tmp_path), str(docx_path),
                ],
                capture_output=True, timeout=60,
            )
        except (subprocess.TimeoutExpired, OSError):
            return None
        pdf_path = tmp_path / "documento.pdf"
        if resultado.returncode != 0 or not pdf_path.exists():
            return None
        return pdf_path.read_bytes()


# --- Marca de agua (sello institucional) sobre cada firma estampada --------
# El sello (Sello_Informatica_Servicios_Generales.svg) se pre-convirtió UNA
# sola vez a PNG semi-transparente, con LibreOffice local, y quedó guardado
# en static/img/formato_permiso/sello_marca_agua.png — así no hace falta
# convertir nada (ni local ni contra Gotenberg) para estampar cada firma.

_marca_agua_cacheada = None


def _obtener_marca_agua():
    """Lee el PNG del sello una sola vez por proceso y lo deja cacheado
    en memoria — no cambia entre solicitudes."""
    global _marca_agua_cacheada
    if _marca_agua_cacheada is None:
        ruta = settings.BASE_DIR / "static" / "img" / "formato_permiso" / "sello_marca_agua.png"
        if ruta.exists():
            _marca_agua_cacheada = ruta.read_bytes()
    return _marca_agua_cacheada


def _agregar_marca_agua(parrafo, doc_pr_id, indent_izquierdo, ancho_columna):
    """Ancla el sello institucional como imagen flotante ARRIBA del
    texto de la firma (behindDoc="1", así no tapa la lectura). `parrafo`
    es el bloque de texto ("Firmado por: ...", etc.) recién insertado
    para esa firma — el desplazamiento vertical negativo empuja el
    sello hacia arriba, para que quede sobre esos datos en vez de
    extenderse hacia abajo, donde se superponía con el bloque del
    siguiente firmante. python-docx no soporta imágenes ancladas
    nativamente, así que se arma el XML del dibujo a mano.

    `positionH relativeFrom="paragraph"` ignora la sangría del párrafo
    en la práctica (LibreOffice lo ancla igual para todas las firmas,
    pegado al margen izquierdo) — por eso se usa `relativeFrom="column"`
    y se replica a mano el mismo desplazamiento (`indent_izquierdo`) que
    ya tiene el párrafo, centrando el sello dentro del ancho real de esa
    columna en vez de dejarlo fijo en el margen."""
    imagen = _obtener_marca_agua()
    if imagen is None:
        return
    rid, _ = parrafo.part.get_or_add_image(io.BytesIO(imagen))
    ancho = int(Inches(1.4))
    alto = int(Inches(1.4))
    offset_x = int(indent_izquierdo) + max(0, (int(ancho_columna) - ancho) // 2)
    offset_y = -(alto + int(Inches(0.1)))
    xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:drawing xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        f'<wp:anchor behindDoc="1" distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        f'locked="0" layoutInCell="1" allowOverlap="1" relativeHeight="{doc_pr_id}">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>{offset_x}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{offset_y}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{ancho}" cy="{alto}"/>'
        '<wp:wrapNone/>'
        f'<wp:docPr id="{doc_pr_id}" name="MarcaAgua{doc_pr_id}"/>'
        '<wp:cNvGraphicFramePr/>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<pic:nvPicPr><pic:cNvPr id="0" name="MarcaAgua{doc_pr_id}"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{ancho}" cy="{alto}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        '</pic:pic></a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>'
    )
    parrafo._p.append(parse_xml(xml))


def renderizar_documento_permiso(solicitud):
    """PDF listo para mostrar/descargar: rellena la plantilla y la
    convierte. Devuelve None si LibreOffice no está disponible o la
    conversión falla, para que la vista pueda avisar en vez de romper."""
    return convertir_docx_a_pdf(rellenar_documento_permiso(solicitud))


def generar_documento_permiso(solicitud):
    """Versión final del formato institucional, guardada una sola vez
    al momento exacto en que la solicitud queda aprobada (mismo
    momento que el comprobante de auditoría)."""
    contenido = renderizar_documento_permiso(solicitud)
    if contenido is None:
        return
    solicitud.documento_permiso.save(
        f"permiso_{solicitud.pk}.pdf", ContentFile(contenido), save=True
    )


def puede_ver_solicitud(user, solicitud):
    """Interesado, cualquiera de sus 3 firmantes, o quien gestione
    solicitudes (RRHH/admin/superuser) — nadie más."""
    funcionario = getattr(user, "funcionario", None)
    es_el_interesado = funcionario is not None and solicitud.funcionario_id == funcionario.id
    es_firmante = funcionario is not None and solicitud.firmas.filter(funcionario_firmante=funcionario).exists()
    return es_el_interesado or es_firmante or puede_gestionar_solicitudes(user)


LOGO_EMAIL = settings.BASE_DIR / "static" / "img" / "formato_permiso" / "logo.png"


def notificar_turno_firma(firma):
    """Avisa por correo a quien le toca firmar un paso, con el logo
    institucional embebido (Content-ID, no un link a imagen externa:
    así se ve igual en cualquier cliente de correo, sin depender de
    que la intranet sea alcanzable desde afuera de la red local). Si
    el funcionario no tiene correo cargado, o el envío falla (SMTP/
    Graph mal configurado, sin conexión, etc.), no interrumpe el
    flujo de firmas: la solicitud sigue firmable desde la app igual,
    solo queda sin notificar."""
    funcionario = firma.funcionario_firmante
    if not funcionario or not funcionario.email:
        return
    url = settings.SITE_URL + reverse("rrhh:detalle_solicitud", args=[firma.solicitud_id])
    contexto = {"firma": firma, "solicitud": firma.solicitud, "url": url}
    texto = render_to_string("rrhh/email_turno_firma.txt", contexto)
    html = render_to_string("rrhh/email_turno_firma.html", contexto)
    try:
        mensaje = EmailMultiAlternatives(
            subject="Tienes una solicitud de permiso pendiente de firma",
            body=texto,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[funcionario.email],
        )
        mensaje.attach_alternative(html, "text/html")
        mensaje.mixed_subtype = "related"
        if LOGO_EMAIL.exists():
            imagen = MIMEImage(LOGO_EMAIL.read_bytes())
            imagen.add_header("Content-ID", "<logo_olivar>")
            imagen.add_header("Content-Disposition", "inline", filename="logo.png")
            mensaje.attach(imagen)
        mensaje.send()
        firma.notificado_en = timezone.now()
        firma.save(update_fields=["notificado_en"])
    except Exception:
        logger.exception(
            "No se pudo enviar el correo de notificación de firma (firma id=%s)", firma.pk
        )


def actualizar_estado_solicitud(solicitud):
    """El estado de la solicitud es derivado de sus firmas, no se
    setea a mano: todas firmadas -> aprobada; alguna rechazada ->
    rechazada; si no, sigue pendiente."""
    firmas = list(solicitud.firmas.select_related("funcionario_firmante").order_by("orden"))
    if any(f.estado == FirmaSolicitud.Estado.RECHAZADO for f in firmas):
        nuevo_estado = SolicitudPermiso.Estado.RECHAZADO
    elif firmas and all(f.estado == FirmaSolicitud.Estado.FIRMADO for f in firmas):
        nuevo_estado = SolicitudPermiso.Estado.APROBADO
    else:
        nuevo_estado = SolicitudPermiso.Estado.PENDIENTE
    if solicitud.estado != nuevo_estado:
        solicitud.estado = nuevo_estado
        solicitud.save(update_fields=["estado"])
        if nuevo_estado == SolicitudPermiso.Estado.APROBADO:
            generar_comprobante_firma(solicitud)
            if tiene_formato_institucional(solicitud):
                generar_documento_permiso(solicitud)
    if nuevo_estado == SolicitudPermiso.Estado.PENDIENTE:
        # El interesado (orden 1) no se notifica por correo: ya queda
        # avisado en pantalla al momento de crear la solicitud. Desde
        # el orden 2 en adelante sí, porque el turno pasa a otra
        # persona que recién ahora puede actuar.
        siguiente = next((f for f in firmas if f.estado == FirmaSolicitud.Estado.PENDIENTE), None)
        if siguiente and siguiente.orden > 1:
            siguiente.solicitud = solicitud
            notificar_turno_firma(siguiente)


def es_su_turno(firma):
    """Una firma solo se puede resolver si todos los pasos anteriores
    (menor `orden`) ya están firmados — no se puede firmar fuera de
    orden."""
    return not firma.solicitud.firmas.filter(orden__lt=firma.orden).exclude(
        estado=FirmaSolicitud.Estado.FIRMADO
    ).exists()


def puede_firmar(user, firma):
    """Solo la persona exacta asignada a esa firma puede firmarla —
    a propósito, sin bypass de superusuario/admin: si un tercero
    pudiera firmar por otro, dejaría de ser una firma electrónica
    válida bajo la Ley 19.799 (no se podría atribuir razonablemente
    el acto a la persona)."""
    return bool(firma.funcionario_firmante) and firma.funcionario_firmante.usuario_id == user.id


def puede_gestionar_solicitudes(user):
    """Determina si se muestra el ítem 'Aprobaciones' en la navegación:
    RRHH/admin/superuser siempre (ven todo), o cualquiera que participe
    como firmante de al menos una solicitud."""
    if not user.is_authenticated:
        return False
    if user.is_superuser or user.rol in ("admin", "rrhh"):
        return True
    funcionario = getattr(user, "funcionario", None)
    return bool(funcionario and funcionario.firmas.exists())


def firmas_visibles_para(user):
    """Bandeja de firmas: RRHH/admin/superuser ven todas (supervisión);
    el resto solo las suyas."""
    qs = FirmaSolicitud.objects.select_related(
        "solicitud", "solicitud__funcionario", "solicitud__funcionario__departamento", "funcionario_firmante"
    )
    if user.is_superuser or user.rol in ("admin", "rrhh"):
        return qs
    funcionario = getattr(user, "funcionario", None)
    if not funcionario:
        return qs.none()
    return qs.filter(funcionario_firmante=funcionario)
